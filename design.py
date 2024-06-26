import sys
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QTextEdit, QLineEdit, QFileDialog, QTableWidget,
    QTableWidgetItem, QAbstractItemView
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QDragEnterEvent, QDropEvent
from docx import Document
import json
from transformers import pipeline
import qtmodern.styles
import qtmodern.windows

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Тест Генератор")
        self.setGeometry(100, 100, 800, 600)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout(self.central_widget)

        self.file_label = QLabel("Перетащите документ сюда или нажмите кнопку для выбора")
        self.file_label.setAcceptDrops(True)
        self.file_label.setStyleSheet("""
            QLabel {
                border: 2px dashed #aaa;
                border-radius: 10px;
                padding: 20px;
                font-size: 16px;
            }
            QLabel:hover {
                border-color: #3a86ff;
            }
        """)
        self.file_label.setAlignment(Qt.AlignCenter)
        self.file_label.setMinimumHeight(100)
        self.file_label.installEventFilter(self)

        self.file_button = QPushButton("Выбрать документ")
        self.file_button.setStyleSheet("""
            QPushButton {
                font-size: 16px;
                padding: 10px 20px;
                border-radius: 5px;
                background-color: #3a86ff;
                color: white;
            }
            QPushButton:hover {
                background-color: #336dbf;
            }
        """)
        self.file_button.clicked.connect(self.open_file_dialog)

        self.topic_label = QLabel("Тема тестирования:")
        self.topic_label.setStyleSheet("font-size: 16px;")
        self.topic_input = QLineEdit()
        self.topic_input.setStyleSheet("""
            QLineEdit {
                font-size: 16px;
                padding: 10px;
                border: 1px solid #ccc;
                border-radius: 5px;
            }
            QLineEdit:focus {
                border-color: #3a86ff;
            }
        """)

        self.num_questions_label = QLabel("Количество вопросов:")
        self.num_questions_label.setStyleSheet("font-size: 16px;")
        self.num_questions_input = QLineEdit()
        self.num_questions_input.setStyleSheet("""
            QLineEdit {
                font-size: 16px;
                padding: 10px;
                border: 1px solid #ccc;
                border-radius: 5px;
            }
            QLineEdit:focus {
                border-color: #3a86ff;
            }
        """)

        self.generate_button = QPushButton("Сформировать вопросы")
        self.generate_button.setStyleSheet("""
            QPushButton {
                font-size: 16px;
                padding: 10px 20px;
                border-radius: 5px;
                background-color: #3a86ff;
                color: white;
            }
            QPushButton:hover {
                background-color: #336dbf;
            }
        """)
        self.generate_button.clicked.connect(self.generate_questions)

        self.questions_output = QTextEdit()
        self.questions_output.setStyleSheet("""
            QTextEdit {
                font-size: 16px;
                padding: 10px;
                border: 1px solid #ccc;
                border-radius: 5px;
            }
            QTextEdit:focus {
                border-color: #3a86ff;
            }
        """)
        self.questions_output.setReadOnly(False)

        self.save_button = QPushButton("Сохранить результаты")
        self.save_button.setStyleSheet("""
            QPushButton {
                font-size: 16px;
                padding: 10px 20px;
                border-radius: 5px;
                background-color: #3a86ff;
                color: white;
            }
            QPushButton:hover {
                background-color: #336dbf;
            }
        """)
        self.save_button.clicked.connect(self.save_results)

        self.view_button = QPushButton("Просмотреть ранее сгенерированные вопросы")
        self.view_button.setStyleSheet("""
            QPushButton {
                font-size: 16px;
                padding: 10px 20px;
                border-radius: 5px;
                background-color: #3a86ff;
                color: white;
            }
            QPushButton:hover {
                background-color: #336dbf;
            }
        """)
        self.view_button.clicked.connect(self.view_previous_questions)

        self.layout.addWidget(self.file_label)
        self.layout.addWidget(self.file_button)
        self.layout.addWidget(self.topic_label)
        self.layout.addWidget(self.topic_input)
        self.layout.addWidget(self.num_questions_label)
        self.layout.addWidget(self.num_questions_input)
        self.layout.addWidget(self.generate_button)
        self.layout.addWidget(self.questions_output)
        self.layout.addWidget(self.save_button)
        self.layout.addWidget(self.view_button)

        self.questions = []
        self.doc_content = ""

    def eventFilter(self, source, event):
        if event.type() == event.DragEnter:
            self.dragEnterEvent(event)
        elif event.type() == event.Drop:
            self.dropEvent(event)
        return super().eventFilter(source, event)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event: QDropEvent):
        urls = event.mimeData().urls()
        if urls:
            file_path = urls[0].toLocalFile()
            if file_path.endswith('.docx'):
                self.load_doc(file_path)

    def open_file_dialog(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Выбрать документ", "", "Документы (*.docx);;Все файлы (*)", options=options)
        if file_name:
            self.load_doc(file_name)

    def load_doc(self, file_path):
        doc = Document(file_path)
        self.doc_content = "\n".join([para.text for para in doc.paragraphs])
        self.file_label.setText(f"Загружен документ: {file_path}")

    def generate_questions(self):
        topic = self.topic_input.text()
        num_questions = int(self.num_questions_input.text())

        generator = pipeline("text2text-generation", model="t5-small")
        input_text = f"topic: {topic}\n\n{self.doc_content}"
        questions = generator(input_text, max_length=50, num_return_sequences=num_questions)

        self.questions = [q['generated_text'] for q in questions]
        self.questions_output.setText("\n".join(self.questions))

    def save_results(self):
        doc = Document()
        doc.add_heading("Сгенерированные вопросы", 0)
        for q in self.questions:
            doc.add_paragraph(q)
        doc.save("generated_questions.docx")

        with open("generated_questions.json", "w", encoding='utf-8') as f:
            json.dump(self.questions, f, ensure_ascii=False, indent=4)

    def view_previous_questions(self):
        self.prev_window = PreviousQuestionsWindow()
        self.prev_window.show()

class PreviousQuestionsWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Ранее сгенерированные вопросы")
        self.setGeometry(100, 100, 800, 600)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout(self.central_widget)

        self.table = QTableWidget()
        self.table.setColumnCount(1)
        self.table.setHorizontalHeaderLabels(["Вопросы"])
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)

        self.back_button = QPushButton("Назад")
        self.back_button.setStyleSheet("""
            QPushButton {
                font-size: 16px;
                padding: 10px 20px;
                border-radius: 5px;
                background-color: #3a86ff;
                color: white;
            }
            QPushButton:hover {
                background-color: #336dbf;
            }
        """)
        self.back_button.clicked.connect(self.close)

        self.layout.addWidget(self.table)
        self.layout.addWidget(self.back_button)

        self.load_previous_questions()

    def load_previous_questions(self):
        try:
            with open("generated_questions.json", "r", encoding='utf-8') as f:
                questions = json.load(f)
                self.table.setRowCount(len(questions))
                for i, question in enumerate(questions):
                    self.table.setItem(i, 0, QTableWidgetItem(question))
        except FileNotFoundError:
            pass

if __name__ == '__main__':
    app = QApplication(sys.argv)
    qtmodern.styles.light(app)
    main_window = MainWindow()
    mw = qtmodern.windows.ModernWindow(main_window)
    mw.show()
    sys.exit(app.exec_())
